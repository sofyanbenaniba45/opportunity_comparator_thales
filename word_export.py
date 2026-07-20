"""
Word report generation (.docx).
"""
import datetime
import io
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
COLOR_HEADER_BG  = "244061"   # Dark blue (hex, no #)
COLOR_HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)  # White
HEX_NEW          = "FFFF00"   # Yellow highlight for new rows
HEX_MODIFIED     = "FFB347"   # Orange highlight for modified cells

FONT_NAME = "Century Gothic"
FONT_SIZE = 10


def _set_cell_bg(cell, hex_color: str):
    """Apply background color to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_run_text_highlight(run, hex_color: str):
    """Apply character-level shading (text highlight) to a run."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _set_cell_text(cell, text: str, bold=False, color=None,
                   font_size=FONT_SIZE, highlight_hex: str = None):
    """Write text into a cell with formatting. Optionally highlight the text only."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text) if text is not None else "")
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if highlight_hex:
        _set_run_text_highlight(run, highlight_hex)


def _normalize_key(val) -> str:
    return str(val).strip().lower()


def _gbu_label(v) -> str:
    s = str(v).strip()
    return "(No GBU/BL)" if s in ("", "nan", "NaT", "None", "<NA>") else s


def _ou_label(v) -> str:
    s = str(v).strip()
    return "(No OU/LE)" if s in ("", "nan", "NaT", "None", "<NA>") else s


def _page_break(doc):
    """Insert a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _apply_section_layout(section):
    """Landscape A4 with narrow margins."""
    section.orientation   = WD_ORIENT.LANDSCAPE
    section.page_width    = Cm(29.7)
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)


def _write_table(doc, columns: list, df_group: pd.DataFrame, diff_map: dict):
    """Add a formatted Word table to the document."""
    n_rows = len(df_group) + 1
    table = doc.add_table(rows=n_rows, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for col_idx, col_name in enumerate(columns):
        cell = hdr_row.cells[col_idx]
        _set_cell_bg(cell, COLOR_HEADER_BG)
        _set_cell_text(cell, col_name, bold=False, color=COLOR_HEADER_FG)

    # Data rows
    for row_idx, (_, row) in enumerate(df_group.iterrows(), start=1):
        status   = row.get("__status__", "unchanged")
        norm_key = _normalize_key(row.get("Opportunity Name", ""))
        changed_cols = diff_map.get(norm_key, set())

        tbl_row = table.rows[row_idx]
        for col_idx, col_name in enumerate(columns):
            cell = tbl_row.cells[col_idx]
            val = row.get(col_name, "")
            if val is None or (isinstance(val, float) and pd.isna(val)):
                val_str = ""
            elif isinstance(val, (datetime.date, datetime.datetime, pd.Timestamp)):
                val_str = val.strftime("%Y-%m-%d")
            else:
                val_str = str(val)

            if status == "new":
                highlight = HEX_NEW
            elif status == "modified" and col_name in changed_cols:
                highlight = HEX_MODIFIED
            else:
                highlight = None

            _set_cell_text(cell, val_str, highlight_hex=highlight)


def _write_opportunity_notes(doc, df_group: pd.DataFrame):
    """
    After a table, write one block per data row:
      Opportunity X : [Opportunity Name] :
      a.Context
      b.Development during the quarter:
      c.Next Steps
    """
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Status/Summary of the projects and the new developments this quarter:")
    r_title.font.name = FONT_NAME
    r_title.font.size = Pt(FONT_SIZE)
    r_title.font.bold = False

    for row_num, (_, row) in enumerate(df_group.iterrows(), start=1):
        opp_name = row.get("Opportunity Name", "")
        opp_name = "" if (opp_name is None or (isinstance(opp_name, float) and pd.isna(opp_name))) else str(opp_name)

        # "1) Opportunity : Name" in bold
        p = doc.add_paragraph()
        run = p.add_run(f"{row_num}) Opportunity : {opp_name}")
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        run.font.bold = True

        for line in ("a. Context:", "b. Development during the quarter:", "c. Next Steps:"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.0)
            r2 = p2.add_run(line)
            r2.font.name = FONT_NAME
            r2.font.size = Pt(FONT_SIZE)
            r2.font.bold = True

        # Blank line between opportunities
        doc.add_paragraph("")


def generate_word(result: dict, gbu_col: str) -> bytes:
    """
    Word report grouped by GBU/BL.
    Heading 1 per GBU/BL, followed by its table.
    Landscape A4, narrow margins, Century Gothic 10pt.
    """
    columns  = result["columns"]
    diff_map = result["diff_map"]

    frames = []
    for status in ("new", "modified", "unchanged"):
        df = result[status]
        if not df.empty:
            frames.append(df)

    doc = Document()
    _apply_section_layout(doc.sections[0])

    if not frames:
        doc.add_paragraph("No data to export.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    df_all = pd.concat(frames, ignore_index=True).copy()

    if gbu_col and gbu_col in df_all.columns:
        df_all["__gbu__"] = [_gbu_label(v) for v in df_all[gbu_col]]
    else:
        df_all["__gbu__"] = "All"

    gbu_values = sorted(df_all["__gbu__"].unique())
    groups = df_all.groupby("__gbu__")

    for gbu in gbu_values:
        df_gbu = groups.get_group(gbu).copy()

        heading = doc.add_heading(str(gbu), level=1)
        if heading.runs:
            heading.runs[0].font.name = FONT_NAME
            heading.runs[0].font.color.rgb = RGBColor(0x24, 0x40, 0x61)

        _write_table(doc, columns, df_gbu, diff_map)
        _write_opportunity_notes(doc, df_gbu)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_word_by_ou(result: dict, ou_col: str, gbu_col: str) -> bytes:
    """
    Word report grouped by Operating Unit / Legal Entity, then by GBU/BL.

    Structure:
      [Page break between each OU/LE]
      Heading 1 : Operating Unit / Legal Entity
        Heading 2 : GBU/BL
          Table   : opportunities for this OU/LE × GBU/BL
    """
    columns  = result["columns"]
    diff_map = result["diff_map"]

    frames = []
    for status in ("new", "modified", "unchanged"):
        df = result[status]
        if not df.empty:
            frames.append(df)

    doc = Document()
    _apply_section_layout(doc.sections[0])

    if not frames:
        doc.add_paragraph("No data to export.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    df_all = pd.concat(frames, ignore_index=True).copy()

    if ou_col and ou_col in df_all.columns:
        df_all["__ou__"] = [_ou_label(v) for v in df_all[ou_col]]
    else:
        df_all["__ou__"] = "All"

    if gbu_col and gbu_col in df_all.columns:
        df_all["__gbu__"] = [_gbu_label(v) for v in df_all[gbu_col]]
    else:
        df_all["__gbu__"] = "All"

    ou_values  = sorted(df_all["__ou__"].unique())
    ou_groups  = df_all.groupby("__ou__")

    for ou_idx, ou in enumerate(ou_values):
        # Page break before each OU/LE except the first
        if ou_idx > 0:
            _page_break(doc)

        # Heading 1: Operating Unit / Legal Entity
        h1 = doc.add_heading(str(ou), level=1)
        if h1.runs:
            h1.runs[0].font.name = FONT_NAME
            h1.runs[0].font.color.rgb = RGBColor(0x24, 0x40, 0x61)

        df_ou = ou_groups.get_group(ou).copy()
        gbu_values = sorted(df_ou["__gbu__"].unique())
        gbu_groups = df_ou.groupby("__gbu__")

        for gbu in gbu_values:
            df_gbu = gbu_groups.get_group(gbu).copy()

            # Heading 2: GBU/BL
            h2 = doc.add_heading(str(gbu), level=2)
            if h2.runs:
                h2.runs[0].font.name = FONT_NAME
                h2.runs[0].font.color.rgb = RGBColor(0x24, 0x40, 0x61)

            _write_table(doc, columns, df_gbu, diff_map)
            _write_opportunity_notes(doc, df_gbu)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
